"""
document_parsing.py

Extracts raw text from insurer-submitted claim documents (PDF or Word)
so it can be handed to the LLM extraction layer. Pure I/O + text
extraction only -- no business logic and no LLM calls live here.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Union

import pdfplumber
from docx import Document as DocxDocument

# Below this many characters of extracted text, a PDF is treated as
# "possibly scanned" -- pdfplumber found little to no embedded text,
# which usually means the page is an image and needs OCR first.
MIN_TEXT_LENGTH_THRESHOLD = 40


@dataclass
class ParsedDocument:
    """Result of extracting text from a single uploaded file."""

    filename: str
    text: str
    possibly_scanned: bool = False
    error: str | None = None

    @property
    def ok(self) -> bool:
        return self.error is None


def parse_pdf(file: Union[str, Path, "object"], filename: str) -> ParsedDocument:
    """Extract text from a PDF using pdfplumber.

    `file` may be a path or a file-like object (e.g. a Streamlit
    UploadedFile), since pdfplumber accepts both.
    """
    try:
        chunks: list[str] = []
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                chunks.append(page_text)
        text = "\n".join(chunks).strip()
    except Exception as exc:  # noqa: BLE001 - surfaced to the UI, not swallowed
        return ParsedDocument(filename=filename, text="", error=f"Failed to read PDF: {exc}")

    possibly_scanned = len(text) < MIN_TEXT_LENGTH_THRESHOLD
    return ParsedDocument(filename=filename, text=text, possibly_scanned=possibly_scanned)


def parse_docx(file: Union[str, Path, "object"], filename: str) -> ParsedDocument:
    """Extract text (paragraphs + table cells) from a Word document."""
    try:
        doc = DocxDocument(file)
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        return ParsedDocument(filename=filename, text="", error=f"Failed to read Word document: {exc}")

    possibly_scanned = len(text) < MIN_TEXT_LENGTH_THRESHOLD
    return ParsedDocument(filename=filename, text=text, possibly_scanned=possibly_scanned)


def parse_document(file: Union[str, Path, "object"], filename: str) -> ParsedDocument:
    """Dispatch to the right parser based on file extension."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file, filename)
    if suffix in (".docx", ".doc"):
        return parse_docx(file, filename)
    return ParsedDocument(
        filename=filename,
        text="",
        error=f"Unsupported file type '{suffix}'. Please upload a PDF or Word (.docx) file.",
    )


def parse_documents(files: list[tuple[Union[str, Path, "object"], str]]) -> list[ParsedDocument]:
    """Parse multiple (file, filename) pairs and return one result per file."""
    return [parse_document(f, name) for f, name in files]


def combine_texts(parsed: list[ParsedDocument]) -> str:
    """Join successfully-parsed documents into one text blob for the LLM,
    labeled by source filename so the model (and a human reviewer) can
    tell which document a fact came from.
    """
    blocks = []
    for doc in parsed:
        if doc.ok and doc.text:
            blocks.append(f"--- Source: {doc.filename} ---\n{doc.text}")
    return "\n\n".join(blocks)
