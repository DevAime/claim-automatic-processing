"""
build_templates.py

One-time build script that generates the two docxtpl source templates
(templates/xl_pla_template.docx and templates/proportional_pla_template.docx)
used by document_generation.py. Mirrors the layout of the two sample
PLA forms supplied by Afro-Asian Reinsurance Brokers.

Run once (or whenever the template layout needs to change):
    python scripts/build_templates.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"

NAVY = RGBColor(0x1B, 0x2A, 0x4A)


def _add_title(doc: Document, heading: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AFRO-ASIAN REINSURANCE BROKERS (KENYA) LTD")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("PRELIMINARY LOSS ADVICE")
    sub_run.bold = True
    sub_run.font.size = Pt(12)

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label.add_run(heading)
    label_run.italic = True
    label_run.font.size = Pt(10)

    doc.add_paragraph("-" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    intro = doc.add_paragraph(
        "We regret to advise you of the loss whose particulars are given below:"
    )
    intro.runs[0].bold = True
    doc.add_paragraph()


def _add_field_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = True
    for i, (label, placeholder) in enumerate(rows):
        label_cell, value_cell = table.rows[i].cells
        label_run = label_cell.paragraphs[0].add_run(label)
        label_run.bold = True
        value_cell.paragraphs[0].add_run(placeholder)


def _add_closing(doc: Document) -> None:
    doc.add_paragraph()
    closing = doc.add_paragraph(
        "We shall keep you informed of further development about this claim. "
        "In the meantime, kindly register the above claim in your books."
    )
    closing.runs[0].bold = True


def build_xl_template() -> None:
    doc = Document()
    _add_title(doc, "FOR PRELIMINARY LOSS ADVICE UNDER EXCESS OF LOSS COVERS")
    rows = [
        ("1. Treaty Type", "{{ treaty_type }}"),
        ("2. Name of Insured", "{{ name_of_insured }}"),
        ("3. Policy No.", "{{ policy_no }}"),
        ("4. Claim No.", "{{ claim_no }}"),
        ("5. Date of Loss", "{{ date_of_loss }}"),
        ("6. Period of Insurance", "{{ period_of_insurance }}"),
        ("7. Sum Insured", "{{ sum_insured }}"),
        ("8. Particulars of Loss", "{{ particulars_of_loss }}"),
        ("9. Description of Risk Covered", "{{ description_of_risk_covered }}"),
        ("10. Estimated Amount of Loss", "{{ estimated_amount_of_loss }}"),
        ("11. Deductible", "{{ deductible }}"),
        (
            "12. Estimated Amount of Loss affected to XL Cover (Layer wise)",
            "{{ estimated_amount_affected_to_xl_cover }}",
        ),
        ("13. Brief Description of Loss", "{{ brief_description_of_loss }}"),
    ]
    _add_field_table(doc, rows)
    _add_closing(doc)
    doc.save(TEMPLATES_DIR / "xl_pla_template.docx")


def build_proportional_template() -> None:
    doc = Document()
    _add_title(doc, "FOR PRELIMINARY LOSS ADVICE UNDER PROPORTIONAL TREATIES")
    rows = [
        ("1. Treaty / U.W. Year", "{{ treaty_uw_year }}"),
        ("2. Name of Insured", "{{ name_of_insured }}"),
        ("3. Policy No.", "{{ policy_no }}"),
        ("4. Claim No.", "{{ claim_no }}"),
        ("5. Date of Loss", "{{ date_of_loss }}"),
        ("6. Period of Insurance", "{{ period_of_insurance }}"),
        ("7. Sum Insured", "{{ sum_insured }}"),
        ("8. Particulars of Loss", "{{ particulars_of_loss }}"),
        ("9. Description of Risk Covered", "{{ description_of_risk_covered }}"),
        ("10. Estimated Amount of Loss", "{{ estimated_amount_of_loss }}"),
        ("11. Retained Loss", "{{ retained_loss }}"),
        ("12. Estimated Amount of Loss affected to Treaty", "{{ estimated_amount_affected_to_treaty }}"),
        ("13. Brief Description of Loss", "{{ brief_description_of_loss }}"),
    ]
    _add_field_table(doc, rows)
    _add_closing(doc)
    doc.save(TEMPLATES_DIR / "proportional_pla_template.docx")


if __name__ == "__main__":
    TEMPLATES_DIR.mkdir(exist_ok=True)
    build_xl_template()
    build_proportional_template()
    print(f"Templates written to {TEMPLATES_DIR}")
